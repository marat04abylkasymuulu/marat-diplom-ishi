#!/usr/bin/env python3
"""
Generate «Глава 2. Практикалык бөлүм» for record-education (документация стилинде).

• Default: thesis/Глава_2_Практикалык_бөлүм_Рекорд_веб.docx — «Глава 2…» шаблонун ачып,
  body гана тазаланат (стил мурунку диплом менен окшош; файл чоң).

• --light: thesis/Глава_2_Практикалык_бөлүм_Рекорд_веб_жеңил.docx — бош Word.

Мазмун: техстек → иштетүү → архитектура + диаграмма + DRF документациясы → моделдер
(түшүнүктөмө + кыска мисалдар) → API/frontend мисалдары → UI (ар бир сүрөт үчүн түшүнүктөмө).
Толук файлдарды көчүрбөйт — түзмө-түз түшүндүрүү жана тандалган код үзүндүлөрү.
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt

URLS: dict[str, str] = {
    "Python": "https://www.python.org/",
    "Django": "https://www.djangoproject.com/",
    "Django_docs": "https://docs.djangoproject.com/en/stable/",
    "DRF": "https://www.django-rest-framework.org/",
    "DRF_api_guide": "https://www.django-rest-framework.org/api-guide/requests/",
    "DRF_routers": "https://www.django-rest-framework.org/api-guide/routers/",
    "DRF_views": "https://www.django-rest-framework.org/api-guide/views/",
    "DRF_serializers": "https://www.django-rest-framework.org/api-guide/serializers/",
    "DRF_browsable": "https://www.django-rest-framework.org/topics/browsable-api/",
    "SimpleJWT": "https://django-rest-framework-simplejwt.readthedocs.io/",
    "SQLite": "https://www.sqlite.org/docs.html",
    "PostgreSQL": "https://www.postgresql.org/docs/",
    "dj_database_url": "https://pypi.org/project/dj-database-url/",
    "python_dotenv": "https://pypi.org/project/python-dotenv/",
    "React": "https://react.dev/",
    "Vite": "https://vite.dev/",
    "React_Router": "https://reactrouter.com/",
    "axios": "https://axios-http.com/",
    "i18next": "https://www.i18next.com/",
    "react_i18next": "https://react.i18next.com/",
    "Tailwind": "https://tailwindcss.com/",
    "Docker": "https://docs.docker.com/",
    "Docker_Compose": "https://docs.docker.com/compose/",
    "Nginx": "https://nginx.org/en/docs/",
    "Gunicorn": "https://docs.gunicorn.org/",
    "WhiteNoise": "http://whitenoise.evans.io/",
    "GitHub": "https://github.com/",
    "GitHub_Actions": "https://docs.github.com/actions",
    "Vitest": "https://vitest.dev/",
    "npm": "https://docs.npmjs.com/",
    "Node_js": "https://nodejs.org/docs/",
    "Pillow": "https://pillow.readthedocs.io/",
    "django_cors_headers": "https://github.com/adamchainz/django-cors-headers",
    "MDN_HTTP": "https://developer.mozilla.org/en-US/docs/Web/HTTP",
    "MDN_CORS": "https://developer.mozilla.org/en-US/docs/Web/HTTP/CORS",
    "JWT": "https://datatracker.ietf.org/doc/html/rfc7519",
    "Git": "https://git-scm.com/doc",
    "repo": "https://github.com/alidin000/record-education",
    "Django_testing": "https://docs.djangoproject.com/en/stable/topics/testing/",
    "Testing_Library": "https://testing-library.com/docs/react-testing-library/intro/",
}


def append_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def plain(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def blank(doc: Document) -> None:
    doc.add_paragraph("")


def add_code_block(doc: Document, code: str, caption: str | None = None) -> None:
    if caption:
        c = doc.add_paragraph(caption)
        for r in c.runs:
            r.italic = True
            r.font.size = Pt(10)
    for line in code.splitlines():
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_together = True
        for r in p.runs:
            r.font.name = "Consolas"
            r.font.size = Pt(9)


def clear_body_keep_sectpr(doc: Document) -> None:
    body = doc.element.body
    for el in list(body):
        if el.tag == qn("w:sectPr"):
            continue
        body.remove(el)


def read_snippet(repo_root: Path, rel: str, line_start: int, line_end: int) -> str:
    """1-based inclusive line numbers; returns empty string on error."""
    try:
        text = (repo_root / rel).read_text(encoding="utf-8")
        lines = text.splitlines()
        return "\n".join(lines[line_start - 1 : line_end])
    except OSError:
        return ""


LINK_TOKEN_ORDER: list[tuple[str, str]] = [
    ("Django REST Framework", URLS["DRF"]),
    ("React Router", URLS["React_Router"]),
    ("django-cors-headers", URLS["django_cors_headers"]),
    ("GitHub Actions", URLS["GitHub_Actions"]),
    ("GitHub", URLS["GitHub"]),
    ("Docker Compose", URLS["Docker_Compose"]),
    ("dj-database-url", URLS["dj_database_url"]),
    ("python-dotenv", URLS["python_dotenv"]),
    ("Node.js", URLS["Node_js"]),
    ("SimpleJWT", URLS["SimpleJWT"]),
    ("PostgreSQL", URLS["PostgreSQL"]),
    ("WhiteNoise", URLS["WhiteNoise"]),
    ("Gunicorn", URLS["Gunicorn"]),
    ("Tailwind CSS", URLS["Tailwind"]),
    ("react-i18next", URLS["react_i18next"]),
    ("SQLite", URLS["SQLite"]),
    ("Vitest", URLS["Vitest"]),
    ("Testing Library", URLS["Testing_Library"]),
    ("Docker", URLS["Docker"]),
    ("Nginx", URLS["Nginx"]),
    ("axios", URLS["axios"]),
    ("CORS", URLS["MDN_CORS"]),
    ("i18next", URLS["i18next"]),
    ("Vite", URLS["Vite"]),
    ("React", URLS["React"]),
    ("Django", URLS["Django_docs"]),
    ("Python", URLS["Python"]),
    ("JWT", URLS["JWT"]),
    ("npm", URLS["npm"]),
    ("Git", URLS["Git"]),
]


def paragraph_with_linked_tokens(doc: Document, sentence: str) -> None:
    p = doc.add_paragraph()
    if not sentence:
        return
    pattern_parts = [re.escape(k) for k, _ in LINK_TOKEN_ORDER]
    pattern = "(" + "|".join(pattern_parts) + ")"
    parts = re.split(pattern, sentence)
    url_by_key = dict(LINK_TOKEN_ORDER)
    for part in parts:
        if not part:
            continue
        if part in url_by_key:
            append_hyperlink(p, url_by_key[part], part)
        else:
            p.add_run(part)


def drf_doc_line(doc: Document) -> None:
    """Окуу үчүн DRF расмий API документациясы."""
    p = doc.add_paragraph()
    p.add_run("DRF API Guide (суроо/жооп, формат): ")
    append_hyperlink(p, URLS["DRF_api_guide"], "Requests")
    p.add_run(" · Routers: ")
    append_hyperlink(p, URLS["DRF_routers"], "Routers")
    p.add_run(" · Views: ")
    append_hyperlink(p, URLS["DRF_views"], "Views")
    p.add_run(" · Serializers: ")
    append_hyperlink(p, URLS["DRF_serializers"], "Serializers")
    p.add_run(" · ")
    append_hyperlink(p, URLS["DRF_browsable"], "Browsable API")


def ui_figure(
    doc: Document,
    num: str,
    title_ky: str,
    purpose_ky: str,
    capture_ky: str,
    note_ky: str | None = None,
) -> None:
    """Сүрөт UI: алдында түшүнүктөмө, ортосунда placeholder, кийин кошумча."""
    plain(doc, f"Сүрөт UI {num}. {title_ky}")
    plain(doc, f"Эмне үчүн керек: {purpose_ky}")
    plain(doc, f"Сүрөттө көрсөтүү керек болгон нерселер: {capture_ky}")
    p = doc.add_paragraph()
    p.add_run(f"[Орун ээлеп — бул жерге скриншот коюңуз: {title_ky}]").italic = True
    if note_ky:
        plain(doc, f"Эскертүү (скриншоттон кийин): {note_ky}")
    blank(doc)


def build_chapter(doc: Document, repo_root: Path) -> None:
    plain(doc, "Глава.2 Практикалык бөлүм")
    blank(doc)
    p0 = doc.add_paragraph()
    p0.add_run("Тема: «Рекорд» окуу борбору үчүн веб-сайт жана админ панели. Бул бөлүм долбоорду "
               "документация форматында түшүндүрөт (орнотуу → архитектура → моделдер → API мисалдары → UI). Репозиторий: ")
    append_hyperlink(p0, URLS["repo"], "github.com/alidin000/record-education")
    p0.add_run(".")
    blank(doc)

    # ----- 2.1 Tech stack -----
    plain(doc, "2.1 Талап кылынган технологиялык стек жана версиялар.")
    blank(doc)
    paragraph_with_linked_tokens(
        doc,
        "Төмөнкү куралдар тандалган: Python, Django, Django REST Framework, SQLite же PostgreSQL, "
        "React, Vite, React Router, axios, i18next, Tailwind CSS, Docker, Gunicorn, Nginx, WhiteNoise, "
        "SimpleJWT, django-cors-headers, Vitest, GitHub Actions.",
    )
    blank(doc)
    plain(
        doc,
        "Версияларды тактоо: backend/requirements.txt жана frontend/package.json файлдарында "
        "көрсөтүлгөн диапазондорду колдонуңуз (мисалы Django 5.1, React 19).",
    )
    blank(doc)

    # ----- 2.2 How to run -----
    plain(doc, "2.2 Долбоорду алуу жана иштетүү (локалдык режим).")
    blank(doc)
    plain(
        doc,
        "Биринчи кезекте долбоорду GitHub аркылуу алуу: «Code → Download ZIP» менен архивди жүктөп, "
        "ачып, record-education түбүнө киребиз. Альтернатива — Git clone.",
    )
    p = doc.add_paragraph()
    p.add_run("Шилтеме: ")
    append_hyperlink(p, URLS["repo"], "record-education")
    blank(doc)
    plain(doc, "2.2.1 Backend.")
    plain(doc, "Виртуалдык чөйрө, көз карандылуулар, миграция, супер-колдонуучу, сервер:")
    add_code_block(
        doc,
        "cd backend\n"
        "python3 -m venv venv && source venv/bin/activate  # Windows: venv\\Scripts\\activate\n"
        "pip install -r requirements.txt\n"
        "cp .env.example .env\n"
        "python manage.py migrate\n"
        "python manage.py createsuperuser\n"
        "python manage.py runserver",
        None,
    )
    plain(doc, "API негизги префикси: http://127.0.0.1:8000/api/ — браузерден DRF Browsable API көрүнүшү мүмкүн.")
    blank(doc)
    drf_doc_line(doc)
    blank(doc)

    plain(doc, "2.2.2 Frontend.")
    plain(doc, "Экинчи терминалда Vite dev-сервер:")
    add_code_block(doc, "cd frontend\nnpm install\nnpm run dev", None)
    paragraph_with_linked_tokens(
        doc,
        "Адатта коомдук сайт http://127.0.0.1:5173 дарегинде ачылат; Vite прокси же VITE_API_URL аркылуу APIга туташат.",
    )
    blank(doc)

    plain(doc, "2.2.3 Docker Compose (кыскача).")
    paragraph_with_linked_tokens(
        doc,
        "Түбүнөн docker compose up --build — PostgreSQL, backend жана frontend бирге козголот (Docker Compose).",
    )
    snippet_dc = read_snippet(repo_root, "docker-compose.yml", 1, 35)
    if snippet_dc:
        add_code_block(
            doc,
            snippet_dc,
            "Мисал: docker-compose.yml (кызматтардын шилтемеси; толук файл репозиторийде).",
        )
    blank(doc)

    # ----- 2.3 Architecture + graph -----
    plain(doc, "2.3 Архитектура жана Frontend–Backend өз ара аракети.")
    blank(doc)
    paragraph_with_linked_tokens(
        doc,
        "Frontend React SPA HTTP JSON менен Django REST Framework endpointтерине кайрылат. "
        "Сервер JSON жооп кайтарат; коопсуздук үчүн CORS жана JWT колдонулат.",
    )
    blank(doc)
    plain(doc, "Логикалык диаграмма (тексттик граф):")
    plain(doc, "  ┌─────────────────────────────────────────┐")
    plain(doc, "  │ Браузер: React (коомдук + /panel админ) │")
    plain(doc, "  └──────────────┬──────────────────────────┘")
    plain(doc, "                 │ HTTPS, JSON")
    plain(doc, "  ┌──────────────▼──────────────────────────┐")
    plain(doc, "  │ Vite dev же Nginx (Docker) статика      │")
    plain(doc, "  └──────────────┬──────────────────────────┘")
    plain(doc, "                 │ /api/... прокси")
    plain(doc, "  ┌──────────────▼──────────────────────────┐")
    plain(doc, "  │ Gunicorn + Django + DRF ViewSet/router  │")
    plain(doc, "  └──────────────┬──────────────────────────┘")
    plain(doc, "                 │ ORM")
    plain(doc, "  ┌──────────────▼──────────────────────────┐")
    plain(doc, "  │ SQLite (локалдык) / PostgreSQL (Docker) │")
    plain(doc, "  └─────────────────────────────────────────┘")
    blank(doc)
    plain(
        doc,
        "Маалымат агымы мисалы (курстар): браузер getCourses() чакырат → GET /api/courses/ → "
        "CourseViewSet queryset сериализацияланат → JSON → React баракчасы тизмеги.",
    )
    blank(doc)
    plain(doc, "Django REST Framework документациясы (API түзүү жана тестирлөө үчүн):")
    drf_doc_line(doc)
    blank(doc)
    p2 = doc.add_paragraph()
    p2.add_run("Жалпы көрсөтмө: ")
    append_hyperlink(p2, URLS["DRF"], "django-rest-framework.org")
    p2.add_run(
        " — ViewSet, permissions, pagination боюнча бөлүмдөр. Браузерден /api/courses/ сыяктуу URL "
        "ачканда Browsable API шаблону пайда болушу мүмкүн (окуучу үчүн ыңгайлуу)."
    )
    blank(doc)

    # ----- 2.4 Data model (explanations + tiny snippets) -----
    plain(doc, "2.4 Маалымат модели жана долбоордун ички структурасы.")
    blank(doc)
    plain(
        doc,
        "Django «app» принциби: courses (курс, категория, расписание), teachers, reviews, news, contacts. "
        "Ар бир маалымат түрү өз моделдеринде; миграция файлдары схеманы версиялайт.",
    )
    blank(doc)
    plain(doc, "2.4.1 Курстар (courses).")
    plain(
        doc,
        "CourseCategory slug менен фильтрлөөгө ыңгайлуу; Course чет тилдик талааларды (ky/ru/en) сактайт; "
        "Schedule курс менен байланыштуу жума күнү боюнча сабактарды көрсөтөт.",
    )
    cm = read_snippet(repo_root, "backend/courses/models.py", 4, 38)
    if cm:
        add_code_block(doc, cm, "Мисал: CourseCategory жана Course (файлдын үзүндүсү).")
    blank(doc)

    plain(doc, "2.4.2 Байланыштар (contacts).")
    plain(
        doc,
        "Branch филиалдарды сактайт; google_maps_embed_url жана two_gis_embed_url — карта iframe үчүн; "
        "SitePromo башкы беттеги акциялык тизме (бир гана логикалык жазуу, pk=1).",
    )
    cm2 = read_snippet(repo_root, "backend/contacts/models.py", 4, 48)
    if cm2:
        add_code_block(doc, cm2, "Мисал: SitePromo жана Branch талаалары (үзүндү).")
    blank(doc)

    plain(doc, "2.4.3 Окутуучулар (teachers).")
    plain(doc, "TeacherViewSet активдүү мугалимдерди гана коомдук APIдан чыгарат (is_active=True).")
    blank(doc)

    # ----- 2.5 Backend API (small examples only) -----
    plain(doc, "2.5 Backend: URL маршрутизациясы жана ViewSet мисалы.")
    blank(doc)
    plain(
        doc,
        "config/urls.py ичинде DefaultRouter катталган ViewSetтерге REST жолдорду автоматтык түзөт. "
        "Кошумча функциялык view'лар: site-promo жана resolve-map-link.",
    )
    urls_snip = read_snippet(repo_root, "backend/config/urls.py", 1, 31)
    if urls_snip:
        add_code_block(doc, urls_snip, "Мисал: backend/config/urls.py (router + кошумча path'тер).")
    blank(doc)

    tv = read_snippet(repo_root, "backend/teachers/views.py", 1, 12)
    if tv:
        add_code_block(
            doc,
            tv,
            "Мисал: ReadOnlyModelViewSet — коомдук окуу үчүн гана окуу (POST жок).",
        )
    blank(doc)

    plain(doc, "Кошумча endpoint: resolve-map-link (share URL → координаталар).")
    rs = read_snippet(repo_root, "backend/contacts/views.py", 37, 53)
    if rs:
        add_code_block(doc, rs, "Мисал: resolve_map_link view (кыска).")
    blank(doc)

    plain(doc, "JWT жана throttling (продакшн vs тест).")
    st = read_snippet(repo_root, "backend/config/settings.py", 133, 156)
    if st:
        add_code_block(doc, st, "Мисал: REST_FRAMEWORK + SIMPLE_JWT (үзүндү).")
    blank(doc)

    # ----- 2.6 Frontend client -----
    plain(doc, "2.6 Frontend: API клиент жана маршруттар.")
    blank(doc)
    plain(
        doc,
        "axios instance базалык URL менен түзүлөт; ар бир ресурс үчүн кичине функциялар "
        "(getCourses, getTeachers ж.б.) кайталанууну азайтат.",
    )
    ap = read_snippet(repo_root, "frontend/src/utils/api.js", 1, 22)
    if ap:
        add_code_block(doc, ap, "Мисал: frontend/src/utils/api.js (башы + бир нече getter).")
    blank(doc)

    plain(doc, "Админ навигациясы navItems массиви менен централизацияланган:")
    ad = read_snippet(repo_root, "frontend/src/admin/components/AdminLayout.jsx", 10, 22)
    if ad:
        add_code_block(doc, ad, "Мисал: AdminLayout.jsx — /panel маршруттары.")
    blank(doc)

    # ----- 2.7 UI usage (each figure explained) -----
    plain(doc, "2.7 Колдонуу: коомдук интерфейс жана админ панели (UI сүрөттөр).")
    blank(doc)
    paragraph_with_linked_tokens(
        doc,
        "Төмөнкү орундарда гана скриншот коюңуз; ар биринин алдында эмне көрсөтүлүшү керектиги жазылган.",
    )
    blank(doc)

    ui_figure(
        doc,
        "7.1",
        "Коомдук сайт — башкы бет",
        "Окуучу жана ата-эне үчүн биринчи таасир; акциялык тизме жана негизги CTA көрүнүшү.",
        "Hero, статистика, акция тизмеси, тез байланыш баскычы.",
        "Скриншоттон кийин кыскача жазып коюңуз: кайсы тилде көрсөтүлгөн (ky/ru/en).",
    )
    ui_figure(
        doc,
        "7.2",
        "Коомдук сайт — байланыштар / филиалдар",
        "Филиалдардын дареги жана карталардын туура embed көрүнүшүн текшерүү.",
        "Филиал картасы, телефон, WhatsApp, карта iframe же placeholder.",
        "Эгер карта жок болсо, ачык эле көрсөтүңүз — анда кандай текст көрүнөт.",
    )
    ui_figure(
        doc,
        "7.3",
        "Админ — кирүү баракчасы (/panel/login)",
        "Кызматкердин JWT менен системага кирүүсү.",
        "Логин формасы, ката билдирүүсү болсо ал да көрүнсүн.",
        "Пароль көрүнбөй тургандуктан, форманын өзүн гана көрсөтүү жетиштүү.",
    )
    ui_figure(
        doc,
        "7.4",
        "Админ — Dashboard",
        "Жаңы кайрылуулар жана күтүүдөгү пикирлердин саны.",
        "Санактар карточкалары, шилтемелер.",
        "Сандар нөл болсо да скриншот алса болот — бул нормалдуу абал.",
    )
    ui_figure(
        doc,
        "7.5",
        "Админ — Курстар / категория",
        "Курс базасын толтуруу процессин документтештирип жатабыз.",
        "Категория тизмеси же «жаңы курс» формасы, баасы жана узактыгы талаалары.",
        "Бир эле скринде форма + тизме болсо, кененирээк көрүнүш тандаңыз.",
    )
    ui_figure(
        doc,
        "7.6",
        "Админ — Филиалдар (карта талаалары)",
        "Embed URL жана share link талааларынын администраторго көрүнүшү.",
        "Google/2GIS URL талаалары, сактоо баскычы.",
        "Share link тесттегенде, серверден кайткан жоопту Postman менен кошо көрсөтүү мүмкүн (кошумча).",
    )
    ui_figure(
        doc,
        "7.7",
        "Админ — Акция / тизме",
        "Башкы беттеги текстти оңдоо.",
        "Кыргызча/орусча/англисче талаалар, күйгүзүү чекбоксу.",
        "Текст кыска болсо да, форманын толук көрүнүшү керек.",
    )
    ui_figure(
        doc,
        "7.8",
        "Админ — Окуучу пикири (модерация)",
        "Жарыялоо же четке кагуу баскычтары.",
        "Тизме, статус, баскычтар.",
        "Бир окуяны тандап, detail көрүнүшүн көрсөтүү сунушталат.",
    )
    blank(doc)

    # ----- 2.8 Tests -----
    plain(doc, "2.8 Сапатты текшерүү (автоматтык жана UI багыттары).")
    blank(doc)
    plain(
        doc,
        "Сапатты текшерүү эки катмарда жүргүзүлөт: (1) сервер жана API — Django "
        "колдонмосунун ички тесттери; (2) клиент — Vitest менен модулдук жана "
        "көрүнүш тесттери. Кошумча түрдө реалдуу браузерде кол менен сценарийлер "
        "(UI «иштейби?») аткарылат — бул E2E автоматташтырууга даярдык.",
    )
    blank(doc)

    plain(doc, "2.8.1 Backend: эмне текшерилет?")
    blank(doc)
    plain(
        doc,
        "Команда: python manage.py test. Тесттер жаңы SQLite маалымат базасында "
        "миграцияларды колдонуп, HTTP статус коддорун жана JSON жоопторун текшерет.",
    )
    blank(doc)
    plain(
        doc,
        "admin_api: JWT алуу жана жаңылоо (токен жок / туура эмес пароль), /api/admin/me/ "
        "жообу, staff эмес колдонуучу үчүн dashboard'ка тыюу, dashboard статистикалары, "
        "site-promo GET/PATCH, студенттик feedback рейтингинин валидациясы (1–5).",
    )
    blank(doc)
    plain(
        doc,
        "courses, teachers, reviews, news, contacts: коомдук жана админ API'лер, "
        "моделдер, уруксаттар (мисалы, окуучу пикирин жарыялоо), байланыш формасы, "
        "филиалдар жана resolve-map-link сыяктуу endpoint'тердин жооптору.",
    )
    blank(doc)
    pdt = doc.add_paragraph()
    pdt.add_run("Django тесттери боюнча расмий маалымат: ")
    append_hyperlink(pdt, URLS["Django_testing"], "Writing and running tests")
    pdt.add_run(".")
    blank(doc)

    plain(doc, "2.8.2 Frontend (Vitest): эмне текшерилет?")
    blank(doc)
    paragraph_with_linked_tokens(
        doc,
        "Команда: npm test же npm test -- --run. Vitest аркылуу utils/api.js модулунда "
        "бардык exportталган функциялардын бар экени текшерилет (getCourses, getSitePromo, "
        "resolveMapLink ж.б.) — бул API катмарынын «контракты» сакталып жатканын билдирет.",
    )
    blank(doc)
    plain(
        doc,
        "App.test.jsx: App компоненти mock API менен рендерден өтөт (башкы бет кулабайт) — "
        "бул минималдуу «дымак» (smoke) тест; толук барактарды текшерүү үчүн кошумча тесттер керек.",
    )
    blank(doc)
    plain(
        doc,
        "BranchMapTabs.test.jsx: карта embed URL'дерин sanitizeMapEmbedSrc аркылуу текшерүү — "
        "Google embed кабыл алынат, browse URL жана http/javascript схемалары четке кагылат; "
        "share link түшүнүгү жана OpenStreetMap embed куруу логикасы текшерилет. "
        "Бул UI коопсуздугу жана туура iframe үчүн маанилүү.",
    )
    blank(doc)
    ptl = doc.add_paragraph()
    ptl.add_run("React Testing Library: ")
    append_hyperlink(ptl, URLS["Testing_Library"], "Introduction")
    ptl.add_run(" — render жана expect колдонулат.")
    blank(doc)

    plain(doc, "2.8.3 UI «иштейби?» — кол менен жана CI.")
    blank(doc)
    plain(
        doc,
        "Автоматтык Vitest тесттери негизинен модулдук логиканы жана биринчи рендерди "
        "каптап жатат; реалдуу колдонуучу агымын (форма толтуруу, меню, карталардын "
        "көрүнүшү) толук имитациялоо үчүн адатта Playwright же Cypress сыяктуу E2E "
        "каралгандар кошулат. Учурда долбоордо мындай E2E жок болсо, төмөнкүдөй "
        "кол менен сценарийлерди жазып, скриншот менен тастыкташ кеңеш: ",
    )
    blank(doc)
    plain(doc, "• Башкы бет: тил которгуч, акция тизмеси, шилтемелер иштейтби.")
    plain(doc, "• Курстар: фильтр, баракча ачылат, 404 жок.")
    plain(doc, "• Байланыш: форма жөнөтүү, ката билдирүүлөрү.")
    plain(doc, "• Админ: логин, курстарды кошуу, филиал картасы, промо сактоо.")
    blank(doc)
    paragraph_with_linked_tokens(
        doc,
        "GitHub Actions: pushдо backend test жана frontend Vitest + production build "
        "орто чөйрөдө кайра-кайра иштетилет — регрессияны эрте байкоого жардам берет.",
    )
    blank(doc)
    blank(doc)
    plain(
        doc,
        "Практикалык бөлүмдө долбоордун орнотулушу, архитектурасы, маалымат модели, API жана "
        "клиент тараптын негизги үлгүлөрү документация стилинде түшүндүрүлдү; UI бөлүгү сүрөттөр менен "
        "толукталат.",
    )
    blank(doc)

    plain(doc, "Колдонулган куралдар жана документация")
    blank(doc)
    bib = [
        ("Python", URLS["Python"]),
        ("Django", URLS["Django_docs"]),
        ("Django REST Framework (баштык)", URLS["DRF"]),
        ("DRF — API Guide: Requests", URLS["DRF_api_guide"]),
        ("DRF — Routers", URLS["DRF_routers"]),
        ("DRF — Views", URLS["DRF_views"]),
        ("DRF — Serializers", URLS["DRF_serializers"]),
        ("DRF — Browsable API", URLS["DRF_browsable"]),
        ("SimpleJWT", URLS["SimpleJWT"]),
        ("React", URLS["React"]),
        ("Vite", URLS["Vite"]),
        ("React Router", URLS["React_Router"]),
        ("axios", URLS["axios"]),
        ("Docker", URLS["Docker"]),
        ("Docker Compose", URLS["Docker_Compose"]),
        ("PostgreSQL", URLS["PostgreSQL"]),
        ("SQLite", URLS["SQLite"]),
        ("Vitest", URLS["Vitest"]),
        ("React Testing Library", URLS["Testing_Library"]),
        ("Django — Testing", URLS["Django_testing"]),
        ("GitHub Actions", URLS["GitHub_Actions"]),
        ("Репозиторий", URLS["repo"]),
    ]
    for label, url in bib:
        p = doc.add_paragraph()
        append_hyperlink(p, url, label)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--light", action="store_true", help="Бош Word (жеңил файл).")
    args = parser.parse_args()

    thesis_dir = Path(__file__).resolve().parent
    repo_root = thesis_dir.parent
    template = repo_root / "Глава 2. Практикалык бөлүм.docx"

    if args.light:
        doc = Document()
        out = thesis_dir / "Глава_2_Практикалык_бөлүм_Рекорд_веб_жеңил.docx"
    elif template.is_file():
        doc = Document(str(template))
        clear_body_keep_sectpr(doc)
        out = thesis_dir / "Глава_2_Практикалык_бөлүм_Рекорд_веб.docx"
    else:
        doc = Document()
        out = thesis_dir / "Глава_2_Практикалык_бөлүм_Рекорд_веб.docx"

    build_chapter(doc, repo_root)
    doc.save(str(out))
    print(f"Wrote {out} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
